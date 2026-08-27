#!/usr/bin/env python3
"""Sync version references in CvarcLogger User Manual.docx and add a row to its
Appendix A "Version History" table for a new release -- the python-docx replacement
for the old hand-rolled Word COM automation (Selection.Find/Replace, Table.Rows.Add),
kept because Word COM proved unreliable on this machine (zombie WINWORD processes,
print-spooler hangs -- see the project_word_com_automation_environment memory).

Only touches "Version X.Y" and "vX.Y" text (the cover page and the section 2.1
title-bar example) -- the bare "X.Y" version numbers in the history table's own rows
are never matched by those patterns, so old entries are left alone.

Usage:
    python update-manual-version.py --old 1.27 --new 1.28 --highlights "What changed in 1.28."
"""
import argparse
import copy
from pathlib import Path

from docx import Document
from docx.table import _Row, Table

DEFAULT_MANUAL_PATH = Path(r"C:\Users\user\Documents\Projects\CvarcLogger\docs\CvarcLogger User Manual.docx")


def find_and_replace_in_paragraph(paragraph, find_text: str, replace_text: str) -> bool:
    """Replace find_text with replace_text in a paragraph, handling text split across
    multiple runs (falls back to merging all run text into the first run)."""
    for run in paragraph.runs:
        if find_text in run.text:
            run.text = run.text.replace(find_text, replace_text)
            return True

    full_text = "".join(r.text for r in paragraph.runs)
    if find_text not in full_text or not paragraph.runs:
        return False
    paragraph.runs[0].text = full_text.replace(find_text, replace_text)
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def find_and_replace(doc: Document, find_text: str, replace_text: str) -> int:
    count = 0
    for paragraph in doc.paragraphs:
        if find_and_replace_in_paragraph(paragraph, find_text, replace_text):
            count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if find_and_replace_in_paragraph(paragraph, find_text, replace_text):
                        count += 1
    return count


def find_history_table(doc: Document) -> Table:
    for table in doc.tables:
        header = [c.text.strip() for c in table.rows[0].cells]
        if header[:2] == ["Version", "Highlights"]:
            return table
    raise RuntimeError('Version History table not found (looked for header "Version"/"Highlights")')


def set_cell_text_preserve_format(cell, text: str) -> None:
    """Like `cell.text = text`, but keeps the existing run's formatting (font size,
    etc.) instead of python-docx's default of replacing it with a fresh, unstyled run."""
    paragraph = cell.paragraphs[0]
    if not paragraph.runs:
        paragraph.add_run(text)
    else:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    for extra_paragraph in cell.paragraphs[1:]:
        extra_paragraph._element.getparent().remove(extra_paragraph._element)


def insert_history_row(table: Table, version: str, highlights: str) -> None:
    """Clone the current top data row (index 1, right after the header) to inherit its
    exact formatting, insert the clone above it, and overwrite its text -- so the new
    row becomes the new index 1 and every existing row shifts down by one."""
    template_tr = table.rows[1]._tr
    new_tr = copy.deepcopy(template_tr)
    template_tr.addprevious(new_tr)

    new_row = _Row(new_tr, table)
    set_cell_text_preserve_format(new_row.cells[0], version)
    set_cell_text_preserve_format(new_row.cells[1], highlights)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("--old", required=True, help="Previous version, e.g. 1.27")
    parser.add_argument("--new", required=True, help="New version, e.g. 1.28")
    parser.add_argument("--highlights", required=True, help="Version History row text for the new version")
    parser.add_argument("--manual", default=str(DEFAULT_MANUAL_PATH), help="Path to the .docx (default: the CvarcLogger manual)")
    args = parser.parse_args()

    manual_path = Path(args.manual)
    doc = Document(manual_path)

    n1 = find_and_replace(doc, f"Version {args.old}", f"Version {args.new}")
    n2 = find_and_replace(doc, f"v{args.old}", f"v{args.new}")
    print(f'Replaced "Version {args.old}" -> "Version {args.new}": {n1} occurrence(s)')
    print(f'Replaced "v{args.old}" -> "v{args.new}": {n2} occurrence(s)')

    table = find_history_table(doc)
    insert_history_row(table, args.new, args.highlights)
    print(f"Inserted history row for {args.new}")

    doc.save(manual_path)
    print(f"Saved: {manual_path}")


if __name__ == "__main__":
    main()
